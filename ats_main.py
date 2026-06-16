import os
import re
import json
import hashlib
import fitz  # PyMuPDF
import spacy
import logging
import shutil
import gc
import psutil
from datetime import datetime
import cv2
import numpy as np
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
import easyocr

# System configuration and knowledge databases
logging.basicConfig(filename='ats_system.log', level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')

# Skills database
KNOWN_SKILLS = [
    # Programming Languages
    "python", "java", "javascript", "typescript", "c++", "c#", "golang", "ruby",
    "kotlin", "swift", "rust", "scala", "php", "dart", "r", "perl", "haskell",
    "lua", "objective-c", "assembly", "matlab", "vba", "groovy", "elixir",
    # Frontend
    "react", "angular", "vue", "svelte", "next.js", "nuxt", "html", "css",
    "tailwind", "bootstrap", "sass", "jquery", "redux", "webpack", "vite",
    "figma", "ui/ux", "material ui", "chakra ui", "ant design", "storybook",
    # Mobile
    "react native", "flutter", "swiftui", "jetpack compose", "ionic",
    "xamarin", "cordova", "expo",
    # Backend
    "node.js", "express", "fastapi", "django", "flask", "spring boot",
    "spring security", "hibernate", "rest api", "graphql", "microservices",
    "hono", "nestjs", "koa", "rails", "laravel", "asp.net",
    # Databases
    "sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch",
    "cassandra", "dynamodb", "firebase", "sqlite", "oracle", "mariadb",
    "neo4j", "couchdb", "supabase", "prisma", "drizzle",
    # Cloud & DevOps
    "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "terraform",
    "ci/cd", "linux", "nginx", "apache tomcat", "ansible", "puppet",
    "chef", "vagrant", "helm", "istio", "prometheus", "grafana",
    "cloudflare", "vercel", "netlify", "heroku", "digitalocean",
    # Data & ML / AI
    "machine learning", "deep learning", "nlp", "data analytics", "data science",
    "pandas", "numpy", "tensorflow", "pytorch", "tableau", "power bi",
    "scikit-learn", "keras", "opencv", "computer vision", "langchain",
    "hugging face", "openai", "llm", "rag", "generative ai",
    "data engineering", "etl",
    # Big Data & Data Engineering
    "spark", "hadoop", "airflow", "dbt", "snowflake", "databricks",
    "bigquery", "redshift", "kafka", "rabbitmq", "flink",
    # Security
    "owasp", "penetration testing", "soc", "siem", "encryption",
    "cybersecurity", "vulnerability assessment",
    # Tools
    "git", "maven", "gradle", "postman", "swagger", "junit", "jest",
    "selenium", "jira", "confluence", "agile", "scrum", "puppeteer",
    "cypress", "playwright", "mocha", "chai",
    # Networking
    "tcp/ip", "dns", "vpn", "load balancing", "cdn",
    # Auth
    "oauth", "jwt", "sso", "ldap", "saml",
    # Design
    "adobe xd", "sketch", "canva", "after effects", "photoshop",
    "illustrator", "blender", "premiere pro",
    # No-code / Low-code
    "zapier", "airtable", "notion", "retool", "power automate",
    # Business / HR / Soft Skills
    "human resource", "recruitment", "management", "event management",
    "marketing", "leadership", "itsm", "project management",
    "content management", "sap", "salesforce", "excel",
    "communication", "teamwork", "problem solving", "critical thinking",
]

# Global cities database
TIER1_CITIES = [
    # India — Major
    "delhi", "new delhi", "mumbai", "bangalore", "bengaluru", "chennai",
    "hyderabad", "kolkata", "pune", "noida", "gurugram", "gurgaon",
    "ghaziabad", "lucknow", "jaipur", "chandigarh", "ahmedabad", "indore",
    "bhopal", "patna", "nagpur", "surat", "coimbatore", "kochi",
    "bhubaneswar", "ranchi", "dehradun", "kanpur", "agra", "varanasi",
    "thiruvananthapuram", "visakhapatnam", "mysore", "mangalore", "hubli",
    "vadodara", "rajkot", "jodhpur", "raipur", "guwahati", "jammu",
    # USA
    "new york", "san francisco", "los angeles", "chicago", "seattle",
    "austin", "boston", "denver", "dallas", "houston", "atlanta",
    "miami", "washington dc", "philadelphia", "phoenix", "san jose",
    "san diego", "portland", "minneapolis", "detroit", "charlotte",
    "raleigh", "nashville", "salt lake city",
    # UK
    "london", "manchester", "birmingham", "edinburgh", "glasgow",
    "leeds", "liverpool", "bristol", "cambridge", "oxford", "cardiff",
    # Europe
    "berlin", "munich", "hamburg", "frankfurt", "paris", "lyon",
    "amsterdam", "rotterdam", "dublin", "zurich", "geneva", "vienna",
    "prague", "warsaw", "barcelona", "madrid", "rome", "milan",
    "lisbon", "stockholm", "oslo", "copenhagen", "helsinki", "brussels",
    # Middle East
    "dubai", "abu dhabi", "riyadh", "jeddah", "doha", "kuwait city",
    "muscat", "manama", "amman",
    # SE Asia & East Asia
    "singapore", "kuala lumpur", "bangkok", "jakarta", "manila",
    "ho chi minh city", "hanoi", "tokyo", "osaka", "seoul", "busan",
    "taipei", "hong kong", "shanghai", "beijing", "shenzhen", "guangzhou",
    # Australia & NZ
    "sydney", "melbourne", "brisbane", "perth", "auckland", "wellington",
    # Canada
    "toronto", "vancouver", "montreal", "ottawa", "calgary",
    # Africa
    "cape town", "johannesburg", "nairobi", "lagos", "cairo", "casablanca",
    # South America
    "sao paulo", "buenos aires", "bogota", "santiago", "lima", "mexico city",
]
TIER2_CITIES = [
    # India — Tier 2
    "rohini", "etawah", "ajmer", "meerut", "allahabad", "prayagraj",
    "mathura", "aligarh", "bareilly", "gorakhpur", "faridabad",
    "moradabad", "jalandhar", "amritsar", "ludhiana", "panipat",
    "karnal", "sonipat", "rohtak", "hisar", "bathinda",
    "udaipur", "kota", "bikaner", "sikar", "bhilwara",
    # US — Secondary
    "santa fe", "berkeley", "palo alto", "mountain view", "sunnyvale",
    "redmond", "bellevue", "irvine", "pasadena", "ann arbor",
    # International Secondary
    "pune", "chandigarh", "nagpur", "indore",
]
ALL_CITIES = TIER1_CITIES + TIER2_CITIES

# Certifications database
KNOWN_CERTIFICATIONS = [
    # AWS
    "aws certified solutions architect", "aws certified developer",
    "aws certified sysops", "aws certified cloud practitioner",
    "aws certified devops engineer", "aws saa", "aws ccp", "aws sap",
    # Azure
    "azure fundamentals", "azure administrator", "azure developer",
    "azure solutions architect", "az-900", "az-104", "az-204", "az-305",
    # GCP
    "google cloud certified", "gcp professional cloud architect",
    "gcp associate cloud engineer", "google cloud engineer",
    # Kubernetes & DevOps
    "cka", "ckad", "cks", "certified kubernetes",
    "docker certified", "hashicorp certified terraform",
    # Security
    "cissp", "ceh", "comptia security+", "comptia network+",
    "comptia a+", "oscp", "ccna", "ccnp",
    # Project Management
    "pmp", "prince2", "capm", "itil", "six sigma",
    "six sigma green belt", "six sigma black belt",
    # Agile
    "csm", "certified scrum master", "psm", "safe agilist",
    "pmi-acp",
    # Data & ML
    "tensorflow developer certificate", "databricks certified",
    "snowflake certification", "microsoft certified data analyst",
    "google data analytics",
    # Testing
    "istqb", "istqb foundation", "istqb advanced",
    # Other
    "oracle certified", "java certified", "salesforce certified",
    "hubspot certified", "google analytics certified",
    "meta certified", "facebook blueprint",
]

# Languages database
KNOWN_LANGUAGES = [
    "english", "hindi", "tamil", "telugu", "kannada", "malayalam",
    "bengali", "marathi", "gujarati", "punjabi", "urdu", "odia",
    "assamese", "sanskrit", "nepali", "sindhi", "konkani", "manipuri",
    "dogri", "bodo", "maithili", "santali",
    "french", "german", "spanish", "portuguese", "italian", "dutch",
    "russian", "chinese", "mandarin", "cantonese", "japanese", "korean",
    "arabic", "turkish", "thai", "vietnamese", "indonesian", "malay",
    "swahili", "persian", "hebrew", "polish", "czech", "swedish",
    "norwegian", "danish", "finnish", "greek", "romanian", "hungarian",
]

LANG_PROFICIENCY = [
    "native", "fluent", "proficient", "professional", "advanced",
    "intermediate", "basic", "beginner", "elementary", "conversational",
    "mother tongue", "native speaker", "working proficiency",
    "full professional", "limited working",
]

# Universities database
KNOWN_UNIVERSITIES = [
    # India — IIT
    "iit bombay", "iit delhi", "iit madras", "iit kanpur", "iit kharagpur",
    "iit roorkee", "iit guwahati", "iit hyderabad", "iit bhu",
    "iit indore", "iit mandi", "iit patna", "iit jodhpur",
    # India — NIT
    "nit trichy", "nit warangal", "nit surathkal", "nit calicut",
    "nit rourkela", "nit allahabad", "nit nagpur", "nit kurukshetra",
    # India — IIIT
    "iiit hyderabad", "iiit delhi", "iiit bangalore", "iiit allahabad",
    # India — Other Top
    "bits pilani", "bits hyderabad", "bits goa", "vit vellore",
    "srm university", "manipal university", "amity university",
    "lpu", "lovely professional university",
    "delhi university", "mumbai university", "calcutta university",
    "anna university", "osmania university", "pune university",
    "jadavpur university", "banaras hindu university", "aligarh muslim university",
    "jamia millia islamia", "jnu", "isb hyderabad", "xlri", "iim",
    "iim ahmedabad", "iim bangalore", "iim calcutta", "iim lucknow",
    "christ university", "symbiosis", "thapar university",
    "pec chandigarh", "dtu", "nsut", "iiitd", "igdtuw",
    # USA
    "mit", "stanford", "harvard", "caltech", "princeton",
    "yale", "columbia", "uc berkeley", "ucla", "carnegie mellon",
    "georgia tech", "university of michigan", "cornell",
    "university of texas", "university of illinois",
    "university of washington", "purdue",
    # UK
    "oxford", "cambridge", "imperial college", "ucl", "lse",
    "university of edinburgh", "university of manchester",
    "king's college london", "university of birmingham",
    # Europe
    "eth zurich", "tu munich", "rwth aachen", "delft university",
    "epfl", "sorbonne", "tu berlin",
    # Others
    "university of toronto", "university of waterloo", "mcgill",
    "national university of singapore", "nanyang technological university",
    "university of melbourne", "university of sydney", "anu",
    "tsinghua university", "peking university", "university of tokyo",
    "kaist", "seoul national university",
]

# Job title keywords
JOB_TITLE_KEYWORDS = [
    "engineer", "developer", "programmer", "architect", "analyst",
    "designer", "manager", "lead", "director", "head", "chief",
    "officer", "consultant", "specialist", "coordinator", "administrator",
    "intern", "trainee", "associate", "senior", "junior", "staff",
    "principal", "fellow", "scientist", "researcher", "devops",
    "sre", "qa", "tester", "sde", "swe", "sse", "mts",
    "vp", "avp", "evp", "svp", "cto", "ceo", "cfo", "coo", "cio",
    "full stack", "frontend", "backend", "data engineer", "ml engineer",
    "product manager", "project manager", "scrum master", "tech lead",
    "team lead", "business analyst", "system administrator",
    "network engineer", "security analyst", "cloud engineer",
    "solution architect", "technical writer", "ui designer",
    "ux researcher", "graphic designer", "content writer",
    "digital marketing", "seo specialist", "hr manager",
    "recruitment", "talent acquisition", "executive",
]

# OCR common breaks and fixes
OCR_CITY_FIXES = {
    "k anpur": "kanpur", "kan pur": "kanpur",
    "ghaz iabad": "ghaziabad", "ghazi abad": "ghaziabad",
    "banga lore": "bangalore", "benga luru": "bengaluru",
    "hyde rabad": "hyderabad", "chen nai": "chennai",
    "mum bai": "mumbai", "kol kata": "kolkata",
    "luck now": "lucknow", "chan digarh": "chandigarh",
    "san fran cisco": "san francisco", "los ange les": "los angeles",
    "washing ton": "washington", "singa pore": "singapore",
}

OCR_SKILL_FIXES = {
    "java script": "javascript", "javas cript": "javascript",
    "type script": "typescript", "types cript": "typescript",
    "node js": "node.js", "node. js": "node.js",
    "react js": "react", "react. js": "react",
    "next js": "next.js", "next. js": "next.js",
    "spring boot": "spring boot",
    "machine learn": "machine learning",
    "deep learn": "deep learning",
    "elastic search": "elasticsearch",
    "tail wind": "tailwind", "tailwind css": "tailwind",
    "boot strap": "bootstrap", "post man": "postman",
    "power bi": "power bi",
    "ui / ux": "ui/ux", "ui/ ux": "ui/ux", "ui /ux": "ui/ux",
    "restful apis": "rest api", "restful api": "rest api",
    "ci / cd": "ci/cd", "ci/ cd": "ci/cd",
    "spring secur": "spring security",
    "data analy": "data analytics",
    "react native": "react native",
    "power auto": "power automate",
    "generative ai": "generative ai", "gen ai": "generative ai",
}


def calculate_optimal_threads(num_files):
    cpu_cores = os.cpu_count() or 2
    available_ram_gb = psutil.virtual_memory().available / (1024 ** 3)
    safe_threads_by_ram = max(1, int(available_ram_gb // 1.2))
    safe_threads_by_cpu = max(1, cpu_cores - 1)
    optimal_threads = min(num_files, safe_threads_by_cpu, safe_threads_by_ram)
    return min(optimal_threads, 5)


print("Loading Enterprise AI & Deep Learning Models...")
global_ocr_reader = easyocr.Reader(['en'], gpu=False)

try:
    global_nlp = spacy.load("en_core_web_md")
except:
    global_nlp = spacy.load("en_core_web_sm")

db_lock = threading.Lock()


# Text and image extraction (PDF, Image, DOCX)

def extract_profile_picture_pdf(doc, filename: str) -> str:
    pic_folder = os.path.join("archive", "profile_pics")
    os.makedirs(pic_folder, exist_ok=True)
    face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
    try:
        for i in range(len(doc)):
            for img in doc[i].get_images(full=True):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                pix = fitz.Pixmap(doc, xref)
                aspect_ratio = pix.width / pix.height if pix.height > 0 else 0
                if pix.width > 80 and pix.height > 80 and 0.4 <= aspect_ratio <= 1.5:
                    nparr = np.frombuffer(image_bytes, np.uint8)
                    img_cv = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                    if img_cv is not None:
                        gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
                        faces = face_cascade.detectMultiScale(gray, scaleFactor=1.1, minNeighbors=4)
                        if len(faces) > 0:
                            safe_name = filename.replace(".pdf", "")
                            img_filepath = os.path.join(pic_folder, f"{safe_name}_profile.{image_ext}")
                            with open(img_filepath, "wb") as f:
                                f.write(image_bytes)
                            return img_filepath
    except Exception:
        pass
    return "No Photo"


def extract_face_from_image(file_path: str, filename: str) -> str:
    pic_folder = os.path.join("archive", "profile_pics")
    os.makedirs(pic_folder, exist_ok=True)
    face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
    try:
        img_cv = cv2.imread(file_path)
        if img_cv is not None:
            gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
            faces = face_cascade.detectMultiScale(gray, scaleFactor=1.1, minNeighbors=5)
            if len(faces) > 0:
                x, y, w, h = faces[0]
                face_img = img_cv[y:y+h, x:x+w]
                safe_name = filename.rsplit(".", 1)[0]
                img_filepath = os.path.join(pic_folder, f"{safe_name}_profile.jpg")
                cv2.imwrite(img_filepath, face_img)
                return img_filepath
    except Exception:
        pass
    return "No Photo"


def get_text_from_file(file_path: str, filename: str) -> tuple:
    """Extract raw text + visual name from PDF, image, or DOCX."""
    global global_ocr_reader
    file_ext = file_path.lower().split('.')[-1]
    raw_text = ""
    visual_name = ""

    NAME_BLACKLIST = [
        "resume", "cv", "curriculum vitae", "contact", "objective",
        "summary", "profile", "professional", "experience", "education",
        "skills", "technical", "work", "projects", "certifications",
        "achievements", "leadership", "awards", "references", "declaration",
        "personal", "hobbies", "interests", "languages", "details",
        "social", "awareness", "campaigns", "cultural", "events",
        "extracurricular", "activities", "positions", "responsibility",
        "engagement", "academic", "volunteer", "other",
    ]

    def extract_name_by_font_size(doc):
        max_size = 0
        best_name = ""
        try:
            for page in doc:
                for block in page.get_text("dict").get("blocks", []):
                    if block.get("type") == 0:
                        for line in block.get("lines", []):
                            line_text = ""
                            line_size = 0
                            for span in line.get("spans", []):
                                text = span.get("text", "").strip()
                                size = span.get("size", 0)
                                if text:
                                    line_text += text + " "
                                    line_size = max(line_size, size)
                            line_text = line_text.strip()
                            # Handle pipe-separated headers: "Name | Gender, Age"
                            if '|' in line_text:
                                line_text = line_text.split('|')[0].strip()
                            alpha_only = re.sub(r'[^A-Za-z\s]', '', line_text).strip()
                            words_check = alpha_only.split()
                            if (line_size > max_size
                                    and len(alpha_only) > 2
                                    and 1 <= len(words_check) <= 4
                                    and not any(b in line_text.lower() for b in NAME_BLACKLIST)):
                                max_size = line_size
                                best_name = alpha_only
            words = best_name.split()
            if len(words) >= 3:
                return f"{words[0]} {words[1]} {words[2]}".title()
            elif len(words) == 2:
                return f"{words[0]} {words[1]}".title()
            elif len(words) == 1:
                return words[0].title()
        except:
            pass
        return ""

    def ocr_name_heuristic(text_list):
        for line in text_list[:8]:
            clean_line = re.sub(r'[^A-Za-z\s]', '', line).strip()
            words = clean_line.split()
            if 1 < len(words) <= 4:
                if not any(b in clean_line.lower() for b in NAME_BLACKLIST):
                    if any(w[0].isupper() for w in words if w):
                        return " ".join(words[:3]).title()
        return ""

    if file_ext == 'pdf':
        doc = fitz.open(file_path)
        full_text = []
        for page in doc:
            blocks = page.get_text("blocks")
            blocks.sort(key=lambda b: (b[1], b[0]))
            for b in blocks:
                full_text.append(re.sub(r'\s+', ' ', b[4].strip()))
        raw_text = " \n ".join(full_text)

        if len(raw_text.strip()) < 50:
            raw_text = ""
            for page in doc:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_data = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
                text_list = global_ocr_reader.readtext(img_data, detail=0)
                raw_text += " \n ".join(text_list) + " \n "
                if not visual_name:
                    visual_name = ocr_name_heuristic(text_list)
        else:
            visual_name = extract_name_by_font_size(doc)

    elif file_ext in ['jpg', 'jpeg', 'png']:
        text_list = global_ocr_reader.readtext(file_path, detail=0)
        raw_text = " \n ".join(text_list)
        visual_name = ocr_name_heuristic(text_list)

    elif file_ext == 'docx':
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            paragraphs.append(cell_text)
            raw_text = " \n ".join(paragraphs)
            # Name heuristic: first non-blank line with 1-4 capitalized words
            for line in paragraphs[:5]:
                clean_line = re.sub(r'[^A-Za-z\s]', '', line).strip()
                words = clean_line.split()
                if 1 < len(words) <= 4:
                    if not any(b in clean_line.lower() for b in NAME_BLACKLIST):
                        if any(w[0].isupper() for w in words if w):
                            visual_name = " ".join(words[:3]).title()
                            break
        except Exception as e:
            logging.error(f"DOCX extraction failed for {filename}: {str(e)}")

    return raw_text, visual_name, file_ext == 'pdf'


# Field-level extraction functions (18 layers)

def _clean_email(email: str) -> str:
    """Clean email: truncate at TLD, trim contaminated local part."""
    email = email.lower().strip()

    # 1. Truncate after TLD (remove trailing text like "comwebsitewww...")
    tld_match = re.search(r'@[a-z0-9.\-]+\.(com|org|net|edu|co\.in|co\.uk|io|dev|app|in)', email)
    if tld_match:
        email = email[:tld_match.end()]

    # 2. Remove linkedin/github contamination
    for domain in ["linkedin.com", "github.com", "vercel.app"]:
        if domain in email:
            email = email.split(domain)[0].rstrip('.')
            if '@' not in email:
                return "Not Found"

    # 3. Clean up local part (before @) — remove prefix contamination
    if '@' in email:
        local, domain = email.split('@', 1)
        if len(local) > 30:
            m = re.search(r'[a-z][a-z0-9._%+\-]{1,29}$', local)
            if m:
                local = m.group(0)
        email = f"{local}@{domain}"

    return email if '@' in email else "Not Found"


def extract_email(raw_text: str) -> str:
    """Extract email with aggressive OCR space-break repair."""
    text = raw_text

    # --- OCR Repair: fix spaces around @ and TLDs ---
    text = re.sub(r'\s*@\s*', '@', text)
    text = re.sub(r'\s*\.\s*com\b', '.com', text)
    text = re.sub(r'\s*\.\s*in\b', '.in', text)
    text = re.sub(r'\s*\.\s*org\b', '.org', text)
    text = re.sub(r'\s*\.\s*net\b', '.net', text)
    text = re.sub(r'\s*\.\s*edu\b', '.edu', text)
    text = re.sub(r'\s*\.\s*co\b', '.co', text)

    # --- Method 1: Labeled email field ("Email:", "E-mail ID:") ---
    labeled = re.search(
        r'(?:e[\-\s]?mail\s*(?:id)?\s*[:;]\s*)([a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})',
        text, re.IGNORECASE)
    if labeled:
        return _clean_email(labeled.group(1))

    # --- Method 2: OCR deep repair — find @ and compress zone around it ---
    at_positions = [i for i, c in enumerate(text) if c == '@']
    for pos in at_positions:
        start = max(0, pos - 25)
        end = min(len(text), pos + 30)
        zone = text[start:end]
        compressed = re.sub(r'\s+', '', zone)
        m = re.search(r'[a-zA-Z0-9._%+\-]{2,40}@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}', compressed)
        if m:
            email = _clean_email(m.group(0))
            if "linkedin.com" not in email and "github.com" not in email and '@' in email:
                return email

    # --- Method 3: Standard regex on repaired text ---
    m = re.search(r'\b[a-zA-Z0-9._%+\-]{2,40}@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}\b', text)
    if m:
        email = _clean_email(m.group(0))
        if '@' in email and '.' in email.split('@')[1]:
            return email

    return "Not Found"


def extract_phone(raw_text: str) -> str:
    """Extract phone with Indian (+91) and international format support."""
    clean = re.sub(r'\s+', ' ', raw_text)

    # Pattern 1: Indian with country code +91-XXXXXXXXXX
    m = re.search(r'(\+?91[\-\s.]?\d{10})\b', clean)
    if m:
        return m.group(1).strip()

    # Pattern 2: + followed by 10-12 digits (handles +9354695861)
    m = re.search(r'(\+\d{10,12})\b', clean)
    if m:
        return m.group(1).strip()

    # Pattern 3: Standalone 10-digit Indian number
    m = re.search(r'(?<!\d)(\d{10})(?!\d)', clean)
    if m:
        return m.group(1)

    # Pattern 4: US/intl format (XXX) XXX-XXXX
    m = re.search(r'(?:\+?\d{1,3}[\-.\s]?)?\(?\d{3}\)?[\-.\s]?\d{3}[\-.\s]?\d{4}', clean)
    if m:
        return m.group(0).strip()

    # Pattern 5: Near phone keywords
    m = re.search(r'(?:phone|mobile|contact|cell|tel)[\s:]*(\+?[\d\-\s]{10,15})', clean, re.IGNORECASE)
    if m:
        return re.sub(r'[^\d+\-]', '', m.group(1)).strip()

    return "Not Found"


def extract_location(raw_text: str) -> str:
    """Extract location using header-first strategy with global city priority."""
    global global_nlp
    text_lower = raw_text.lower()

    # Apply OCR city fixes
    for bad, good in OCR_CITY_FIXES.items():
        text_lower = text_lower.replace(bad, good)

    header = text_lower[:500]

    # Strategy 1: "City, Country" pattern in header
    country_pattern = r'(?:india|usa|us|uk|united states|united kingdom|canada|australia|germany|france|singapore|uae|dubai)'
    pattern = r'\b(' + '|'.join(re.escape(c) for c in ALL_CITIES) + r')\s*[,\s]+\s*' + country_pattern
    m = re.search(pattern, header)
    if m:
        return m.group(1).title()

    # Strategy 2: Labeled "Address:" or "Location:" field — prefer Tier1
    addr_match = re.search(r'(?:postal\s+)?(?:address|location|city)\s*[:]\s*(.{5,80})', text_lower)
    if addr_match:
        addr_text = addr_match.group(1)
        for city in TIER1_CITIES:
            if city in addr_text:
                return city.title()
        for city in TIER2_CITIES:
            if city in addr_text:
                return city.title()

    # Strategy 3: Known cities in header — Tier1 first
    for city in TIER1_CITIES:
        if city in header:
            return city.title()
    for city in TIER2_CITIES:
        if city in header:
            return city.title()

    # Strategy 4: NLP GPE from header
    try:
        clean_header = re.sub(r'\s+', ' ', raw_text[:500])
        nlp_doc = global_nlp(clean_header)
        gpe_blacklist = {"india", "us", "usa", "uk", "city", "state", "md", "ma",
                         "ba", "ms", "phd", "java", "python", "bsc", "msc", "mca",
                         "linkedin", "github", "avenue", "redmond", "street",
                         "road", "lane", "nagar", "sector", "block", "phase"}
        for ent in nlp_doc.ents:
            if ent.label_ == "GPE":
                loc = ent.text.strip()
                if loc.lower() not in gpe_blacklist and len(loc) > 2:
                    return loc.title()
    except:
        pass

    # Strategy 5: Full text scan (last resort)
    for city in TIER1_CITIES:
        if city in text_lower:
            return city.title()
    for city in TIER2_CITIES:
        if city in text_lower:
            return city.title()

    return "Not Found"


def extract_education(text_lower: str) -> str:
    """Extract highest education level from text."""
    alpha = re.sub(r'[^a-z0-9\s]', ' ', text_lower)
    alpha = re.sub(r'\s+', ' ', alpha)

    if re.search(r'\b(ph\s*d|phd|doctorate|doctor of)\b', alpha):
        return "Doctorate / PhD"
    if re.search(r'\b(master|masters|m\s?tech|mtech|m\s?s\s?c|msc|m\s?c\s?a|mca|m\s?b\s?a|mba|m\s?com|mcom|p\s?g\s?d\s?m|pgdm)\b', alpha):
        return "Master's Degree / PG"
    if re.search(r'\b(bachelor|bachelors|b\s?tech|btech|b\s?s\s?c|bsc|b\s?c\s?a|bca|b\s?b\s?a|bba|b\s?com|bcom|b\s?e\b)\b', alpha):
        return "Bachelor's Degree"
    if re.search(r'\b(diploma|polytechnic)\b', alpha):
        return "Diploma"
    if re.search(r'\b(degree|graduate|graduation|university|engineering)\b', alpha):
        return "Degree (Unspecified)"
    return "Not Found"


def extract_experience(text_lower: str) -> list:
    """Extract experience: explicit years → intern check → date-range calc → fresher."""
    experience_arr = []

    # Method 1: Explicit "X years/yrs of experience"
    m = re.search(r'(\d+)\s*\+?\s*(?:years?|yrs?)(?:\s+of\s+experience)?', text_lower)
    if m and int(m.group(1)) > 0:
        experience_arr.append(f"{m.group(1)} Years")
        return experience_arr

    # Method 2: Explicit "X months"
    m = re.search(r'(\d+)\s*(?:months?|mos?)(?:\s+of\s+experience)?', text_lower)
    if m:
        months = int(m.group(1))
        if months >= 12:
            experience_arr.append(f"{months // 12} Years")
        elif months > 0:
            experience_arr.append(f"{months} Months")
        return experience_arr

    # Method 3: Intern/trainee keyword → short-circuit
    if re.search(r'\b(intern\b|internship|trainee|apprentice)', text_lower):
        experience_arr.append("Internship / < 1 Year")
        return experience_arr

    # Method 4: Calculate from date ranges in WORK EXPERIENCE section only
    work_section = text_lower
    work_start = re.search(r'\b(work\s+experience|experience|employment|career)\b', text_lower)
    work_end = re.search(r'\b(education|certif|skill|project|award|achievement)\b',
                         text_lower[work_start.end():] if work_start else "")
    if work_start:
        end_pos = work_start.end() + work_end.start() if work_end else len(text_lower)
        work_section = text_lower[work_start.start():end_pos]

    date_ranges = re.findall(
        r'((?:19|20)\d{2})\s*[-–—]+\s*(present|current|now|(?:19|20)\d{2})', work_section)

    if date_ranges:
        current_year = datetime.now().year
        max_calc = 0
        for start_str, end_str in date_ranges:
            try:
                start_yr = int(start_str)
                end_yr = current_year if end_str in ['present', 'current', 'now'] else int(end_str)
                if end_yr >= start_yr and (end_yr - start_yr) <= 20:
                    max_calc = max(max_calc, end_yr - start_yr)
            except:
                pass
        if max_calc > 0:
            experience_arr.append(f"~{max_calc} Years (Calculated)")
            return experience_arr

    experience_arr.append("Fresher")
    return experience_arr


def extract_skills(raw_text: str, custom_skills: list = None) -> list:
    """Extract skills with OCR-aware preprocessing and dynamic expected skills."""
    text = raw_text.lower()

    # Apply OCR skill fixes
    for bad, good in OCR_SKILL_FIXES.items():
        text = text.replace(bad, good)

    search_skills = list(KNOWN_SKILLS)
    if custom_skills:
        for s in custom_skills:
            clean_s = str(s).lower().strip()
            if clean_s and clean_s not in search_skills:
                search_skills.append(clean_s)

    found = set()
    for skill in search_skills:
        try:
            if re.search(r'\b' + re.escape(skill) + r'\b', text):
                found.add(skill)
        except:
            if skill in text:
                found.add(skill)

    # Normalize duplicates
    if "elasticsearch" in found:
        found.discard("elastic search")

    return sorted(list(found))


# LinkedIn, GitHub, and Portfolio URLs

def extract_linkedin_github(raw_text: str) -> dict:
    """Extract LinkedIn, GitHub, and portfolio/website URLs."""
    text = raw_text
    # OCR repair: fix broken URLs
    text = re.sub(r'linked\s*in\s*\.\s*com', 'linkedin.com', text, flags=re.IGNORECASE)
    text = re.sub(r'git\s*hub\s*\.\s*com', 'github.com', text, flags=re.IGNORECASE)
    text = re.sub(r'https?\s*:\s*/\s*/', 'https://', text)

    result = {"linkedin": "", "github": "", "portfolio": ""}

    # LinkedIn
    m = re.search(r'(?:https?://)?(?:www\.)?linkedin\.com/in/([a-zA-Z0-9\-_%]+)', text, re.IGNORECASE)
    if m:
        result["linkedin"] = f"linkedin.com/in/{m.group(1).rstrip('/')}"

    # GitHub
    m = re.search(r'(?:https?://)?(?:www\.)?github\.com/([a-zA-Z0-9\-_]+)', text, re.IGNORECASE)
    if m:
        username = m.group(1).lower()
        # Filter out common non-user GitHub paths
        if username not in ['topics', 'explore', 'settings', 'notifications', 'about',
                            'pricing', 'features', 'enterprise', 'marketplace']:
            result["github"] = f"github.com/{m.group(1)}"

    # Portfolio / Personal Website
    portfolio_keywords = ['portfolio', 'website', 'blog', 'personal site', 'web']
    # Look for URLs near portfolio keywords
    for kw in portfolio_keywords:
        pattern = rf'{re.escape(kw)}[:\s]*(?:https?://)?([a-zA-Z0-9\-]+\.[a-zA-Z0-9.\-/]+)'
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            url = m.group(1).strip().rstrip('.')
            if 'linkedin.com' not in url and 'github.com' not in url:
                result["portfolio"] = url
                break

    # Fallback: Any .dev, .io, .me, .portfolio domain
    if not result["portfolio"]:
        m = re.search(r'(?:https?://)?([a-zA-Z0-9][a-zA-Z0-9\-]*\.(dev|io|me|design|tech|codes|website|site|app)(?:/[^\s]*)?)', text)
        if m:
            url = m.group(1).strip().rstrip('.')
            if 'linkedin.com' not in url and 'github.com' not in url and 'easyocr' not in url:
                result["portfolio"] = url

    return result


# Professional Certifications

def extract_certifications(raw_text: str) -> list:
    """Extract professional certifications from resume."""
    text_lower = raw_text.lower()
    found = []

    # Method 1: Match against known certifications database
    for cert in KNOWN_CERTIFICATIONS:
        try:
            if re.search(r'\b' + re.escape(cert) + r'\b', text_lower):
                found.append(cert.title())
        except:
            if cert in text_lower:
                found.append(cert.title())

    # Method 2: Pattern matching for "Certified [X]"
    cert_patterns = re.findall(
        r'(?:certified|certification|certificate)\s+(?:in\s+)?([A-Za-z\s\-]+?)(?:\s*[\-–|,\n]|\s*$)',
        text_lower)
    for match in cert_patterns:
        clean = match.strip().title()
        if 3 < len(clean) < 60 and clean not in found:
            found.append(clean)

    # Method 3: Section-header based extraction
    cert_section = re.search(
        r'(?:certif(?:ication|icate)s?|licenses?|credentials?)\s*[:\-]?\s*\n?(.*?)(?:\n\s*\n|$)',
        text_lower, re.DOTALL)
    if cert_section:
        lines = cert_section.group(1).split('\n')
        for line in lines[:10]:
            line = line.strip()
            if 5 < len(line) < 80 and not re.match(r'^[\-•●○►▸]?\s*$', line):
                clean = re.sub(r'^[\-•●○►▸]\s*', '', line).strip().title()
                if clean and clean not in found:
                    found.append(clean)

    # Deduplicate (case insensitive)
    seen = set()
    deduped = []
    for c in found:
        key = c.lower()
        if key not in seen:
            seen.add(key)
            deduped.append(c)

    return deduped[:15]  # Cap at 15 certifications


# Languages Spoken

def extract_languages(raw_text: str) -> list:
    """Extract spoken languages with proficiency levels."""
    text_lower = raw_text.lower()
    found = []

    # Try to isolate the languages section
    lang_section = ""
    lang_match = re.search(
        r'(?:languages?\s*(?:known|spoken|proficiency|skills)?)\s*[:\-]?\s*\n?(.*?)(?:\n\s*\n|\b(?:hobbies?|interests?|declaration|references?|skills?|projects?)\b)',
        text_lower, re.DOTALL)
    if lang_match:
        lang_section = lang_match.group(1)

    search_text = lang_section if lang_section else text_lower

    for lang in KNOWN_LANGUAGES:
        if re.search(r'\b' + re.escape(lang) + r'\b', search_text):
            # Try to find proficiency level nearby
            proficiency = ""
            for prof in LANG_PROFICIENCY:
                pattern = rf'{re.escape(lang)}\s*[\-:–(]?\s*{re.escape(prof)}'
                pattern2 = rf'{re.escape(prof)}\s*[\-:–)]?\s*{re.escape(lang)}'
                if re.search(pattern, search_text) or re.search(pattern2, search_text):
                    proficiency = prof.title()
                    break

            entry = lang.title()
            if proficiency:
                entry += f" ({proficiency})"
            found.append(entry)

    return found[:10]  # Cap at 10 languages


# Professional Summary / Objective

def extract_summary(raw_text: str) -> str:
    """Extract the professional summary or objective section."""
    text = raw_text

    # Look for common summary section headers
    headers = [
        r'(?:professional\s+)?summary',
        r'(?:career\s+)?objective',
        r'about\s+me',
        r'profile\s+summary',
        r'(?:professional\s+)?profile',
        r'executive\s+summary',
        r'personal\s+statement',
    ]

    for header in headers:
        pattern = rf'(?:^|\n)\s*{header}\s*[:\-]?\s*\n?(.*?)(?:\n\s*\n|\b(?:experience|education|skills?|technical|work|projects?|certif|employment)\b)'
        m = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if m:
            summary = m.group(1).strip()
            # Clean up
            summary = re.sub(r'\s+', ' ', summary)
            summary = re.sub(r'<[^>]+>', '', summary)  # Remove HTML if any
            if len(summary) > 20:
                return summary[:500]  # Cap at 500 characters

    return ""


# University / Institution Name

def extract_university(raw_text: str) -> str:
    """Extract the university/institution name from resume."""
    text_lower = raw_text.lower()

    # Method 1: Match from known universities database
    for uni in KNOWN_UNIVERSITIES:
        if re.search(r'\b' + re.escape(uni) + r'\b', text_lower):
            return uni.title()

    # Method 2: Pattern matching — "from [University]" or "at [University]"
    uni_patterns = [
        r'(?:from|at)\s+([A-Z][A-Za-z\s]+(?:University|Institute|College|School|Academy))',
        r'([A-Z][A-Za-z\s]+(?:University|Institute|College|School|Academy))\s*[,\-–]',
        r'(?:University|Institute|College)\s+of\s+([A-Z][A-Za-z\s]+)',
    ]
    for pattern in uni_patterns:
        m = re.search(pattern, raw_text)
        if m:
            name = m.group(1).strip()
            if 3 < len(name) < 80:
                return name.title()

    # Method 3: spaCy ORG entity within education section
    global global_nlp
    edu_match = re.search(r'(?:education|academic|qualification)\b(.*?)(?:experience|skills|projects|certif)',
                          text_lower, re.DOTALL)
    if edu_match:
        edu_text = raw_text[edu_match.start():edu_match.end()][:500]
        try:
            nlp_doc = global_nlp(edu_text)
            for ent in nlp_doc.ents:
                if ent.label_ == "ORG":
                    org = ent.text.strip()
                    if len(org) > 3 and any(kw in org.lower() for kw in
                        ['university', 'institute', 'college', 'school', 'iit', 'nit',
                         'iiit', 'bits', 'vit', 'srm', 'academy']):
                        return org.title()
        except:
            pass

    return ""


# Graduation Year

def extract_graduation_year(raw_text: str) -> str:
    """Extract the most recent graduation year from the education section."""
    text_lower = raw_text.lower()

    # Isolate education section
    edu_text = text_lower
    edu_start = re.search(r'\b(?:education|academic|qualification)\b', text_lower)
    edu_end = re.search(r'\b(?:experience|skills|projects|certif|work)\b',
                        text_lower[edu_start.end():] if edu_start else "")
    if edu_start:
        end_pos = edu_start.end() + edu_end.start() if edu_end else len(text_lower)
        edu_text = text_lower[edu_start.start():end_pos]

    # Find 4-digit years in education context
    years = re.findall(r'\b(20[0-3]\d|19[89]\d)\b', edu_text)
    if years:
        # Return the most recent year
        int_years = [int(y) for y in years]
        current_year = datetime.now().year
        valid_years = [y for y in int_years if y <= current_year + 1]
        if valid_years:
            return str(max(valid_years))

    return ""


# Company Names and Job Titles

def extract_work_history(raw_text: str) -> list:
    """Extract structured work history: company, title, duration per role."""
    global global_nlp
    text_lower = raw_text.lower()
    work_history = []

    # Isolate work experience section
    work_text = raw_text
    work_start = re.search(r'(?:work\s+experience|professional\s+experience|experience|employment\s+history|career)',
                           text_lower)
    work_end = re.search(r'\b(?:education|certif|skill|project|award|achievement|academic)\b',
                         text_lower[work_start.end():] if work_start else "")
    if work_start:
        end_pos = work_start.end() + work_end.start() if work_end else len(text_lower)
        work_text = raw_text[work_start.start():end_pos]

    if len(work_text.strip()) < 30:
        return work_history

    # Find date ranges to split into role blocks
    date_pattern = r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|(?:19|20)\d{2})\s*[-–—]+\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|(?:19|20)\d{2}|[Pp]resent|[Cc]urrent|[Nn]ow)'
    date_matches = list(re.finditer(date_pattern, work_text))

    if not date_matches:
        return work_history

    # Use spaCy to find ORG entities (company names)
    try:
        nlp_doc = global_nlp(work_text[:2000])
        org_entities = [ent.text.strip() for ent in nlp_doc.ents if ent.label_ == "ORG"]
    except:
        org_entities = []

    for i, dm in enumerate(date_matches[:5]):  # Cap at 5 roles
        duration = f"{dm.group(1)} – {dm.group(2)}"

        # Look at text around this date match for title and company
        start = max(0, dm.start() - 200)
        end = min(len(work_text), dm.end() + 100)
        context = work_text[start:end]
        context_lower = context.lower()

        # Find job title
        title = ""
        for jt in JOB_TITLE_KEYWORDS:
            if jt in context_lower:
                # Try to extract the full title line
                for line in context.split('\n'):
                    if jt in line.lower() and len(line.strip()) < 80:
                        title = re.sub(r'\s+', ' ', line.strip()).title()
                        break
                if title:
                    break

        # Find company name
        company = ""
        for org in org_entities:
            if org.lower() in context_lower and len(org) > 2:
                company = org.title()
                break

        if title or company:
            entry = {}
            if company:
                entry["company"] = company
            if title:
                entry["title"] = title[:80]
            entry["duration"] = duration
            work_history.append(entry)

    return work_history[:5]  # Cap at 5 entries


# Duplicate Detection

def generate_candidate_fingerprint(name: str, email: str, phone: str) -> str:
    """Generate a fingerprint hash for duplicate detection."""
    normalized = f"{name.lower().strip()}|{email.lower().strip()}|{phone.strip()}"
    return hashlib.md5(normalized.encode()).hexdigest()


# Main resume processor

def process_single_resume(file_path: str, filename: str, expected_skills: list = None) -> dict:
    global global_nlp
    try:
        if os.path.getsize(file_path) == 0:
            return {"status": "failed", "filename": filename, "error": "File is empty"}

        raw_text, visual_name, is_pdf = get_text_from_file(file_path, filename)

        if not raw_text.strip():
            return {"status": "failed", "filename": filename, "error": "No text found"}

        # Extract fields
        email = extract_email(raw_text)
        phone = extract_phone(raw_text)
        location = extract_location(raw_text)

        clean_text = re.sub(r'\s+', ' ', raw_text)
        text_lower = clean_text.lower()

        education = extract_education(text_lower)
        experience = extract_experience(text_lower)
        skills = extract_skills(raw_text, expected_skills)

        # New layers
        urls = extract_linkedin_github(raw_text)
        certifications = extract_certifications(raw_text)
        languages = extract_languages(raw_text)
        summary = extract_summary(raw_text)
        university = extract_university(raw_text)
        grad_year = extract_graduation_year(raw_text)
        work_history = extract_work_history(raw_text)

        # Profile picture
        file_ext = file_path.lower().split('.')[-1]
        if file_ext == 'pdf':
            doc = fitz.open(file_path)
            profile_pic_path = extract_profile_picture_pdf(doc, filename)
            doc.close()
        elif file_ext in ['jpg', 'jpeg', 'png']:
            profile_pic_path = extract_face_from_image(file_path, filename)
        else:
            profile_pic_path = "No Photo"

        # Name extraction strategy
        if not visual_name or len(visual_name.strip()) < 3:
            nlp_doc = global_nlp(clean_text[:500])
            persons = [ent.text.strip().title() for ent in nlp_doc.ents if ent.label_ == "PERSON"]
            if persons:
                visual_name = persons[0]
            elif email != "Not Found":
                visual_name = re.sub(r'[0-9._\-]', ' ', email.split('@')[0]).strip().title()
            else:
                visual_name = "Unknown Candidate"

        vw = visual_name.split()
        if len(vw) > 3:
            visual_name = f"{vw[0]} {vw[1]} {vw[2]}"

        # Generate duplicate fingerprint
        fingerprint = generate_candidate_fingerprint(visual_name, email, phone)

        return {
            "status": "success",
            "filename": filename,
            "data": {
                "name": visual_name,
                "email": email,
                "phone": phone,
                "location": location,
                "education": education,
                "experience": experience,
                "skills": skills,
                "profile_pic": profile_pic_path,
                "linkedin": urls.get("linkedin", ""),
                "github": urls.get("github", ""),
                "portfolio": urls.get("portfolio", ""),
                "certifications": certifications,
                "languages": languages,
                "summary": summary,
                "university": university,
                "grad_year": grad_year,
                "work_history": work_history,
                "fingerprint": fingerprint,
                "status": "Processed"
            }
        }
    except Exception as e:
        logging.error(f"Exception processing {filename}: {str(e)}")
        return {"status": "failed", "filename": filename, "error": str(e)}
    finally:
        gc.collect()


# Pipeline orchestrator and multi-factor scoring engine

def calculate_multi_factor_score(candidate_data: dict, expected_skills: list,
                                  jd_skills: set) -> int:
    """
    Multi-factor weighted scoring:
    Skills (50%) + Experience (20%) + Education (15%) + Certifications (15%)
    """
    # Skill Match (50%)
    cand_skills = candidate_data.get('skills', [])
    if isinstance(cand_skills, str):
        try:
            cand_skills = json.loads(cand_skills)
        except:
            cand_skills = []

    cand_set = set(str(s).lower().strip() for s in cand_skills)
    manual_exp_set = set(str(s).lower().strip() for s in expected_skills if str(s).strip())
    exp_set = manual_exp_set.union(jd_skills)

    if len(exp_set) > 0:
        matched = cand_set.intersection(exp_set)
        skill_score = (len(matched) / len(exp_set)) * 100
    else:
        skill_score = min(len(cand_set) * 5, 85)

    # Experience Fit (20%)
    exp_raw = candidate_data.get('experience', ['Fresher'])
    if isinstance(exp_raw, list):
        exp_str = exp_raw[0] if exp_raw else "Fresher"
    else:
        exp_str = str(exp_raw)

    exp_score = 30  # Default for Fresher
    years_match = re.search(r'(\d+)', exp_str)
    if years_match:
        years = int(years_match.group(1))
        if years >= 8:
            exp_score = 100
        elif years >= 5:
            exp_score = 85
        elif years >= 3:
            exp_score = 70
        elif years >= 1:
            exp_score = 55
        else:
            exp_score = 40
    elif 'internship' in exp_str.lower():
        exp_score = 35

    # Education Fit (15%)
    edu = candidate_data.get('education', 'Not Found')
    edu_score_map = {
        "Doctorate / PhD": 100,
        "Master's Degree / PG": 85,
        "Bachelor's Degree": 70,
        "Diploma": 50,
        "Degree (Unspecified)": 55,
    }
    edu_score = edu_score_map.get(edu, 20)

    # Certification Bonus (15%)
    certs = candidate_data.get('certifications', [])
    cert_score = min(len(certs) * 25, 100)  # Each cert = 25 points, capped at 100

    # ── Weighted combination ──
    final = (skill_score * 0.50) + (exp_score * 0.20) + (edu_score * 0.15) + (cert_score * 0.15)
    return int(min(max(final, 0), 100))


def run_ats_pipeline(expected_skills=None, job_id=None, job_description=None):
    if expected_skills is None:
        expected_skills = []
    TARGET_FOLDER = "resumes_to_process"
    PROCESSED_FOLDER = "archive/processed"
    FAILED_FOLDER = "archive/failed"
    PICS_FOLDER = "archive/profile_pics"
    JSON_OUTPUT_FILE = "candidates_data.json"

    for folder in [TARGET_FOLDER, PROCESSED_FOLDER, FAILED_FOLDER, PICS_FOLDER]:
        os.makedirs(folder, exist_ok=True)

    files = [f for f in os.listdir(TARGET_FOLDER)
             if f.lower().endswith(('.pdf', '.png', '.jpg', '.jpeg', '.docx'))]
    if not files:
        print(f"\n[!] Folder '{TARGET_FOLDER}' is empty.")
        return

    optimal_threads = calculate_optimal_threads(len(files))
    print(f"\n [Start] Enterprise Pipeline: Processing {len(files)} files with {optimal_threads} threads...")

    # Extract skills from JD if provided (done once, shared across all resumes)
    jd_skills = set()
    if job_description:
        clean_jd = re.sub(r'<[^>]+>', ' ', job_description).lower()
        for k_skill in KNOWN_SKILLS:
            if re.search(r'\b' + re.escape(k_skill) + r'\b', clean_jd):
                jd_skills.add(k_skill)

    extracted_candidates_data = []

    with ThreadPoolExecutor(max_workers=optimal_threads) as executor:
        futures = {executor.submit(process_single_resume, os.path.join(TARGET_FOLDER, f), f, expected_skills): f for f in files}

        for future in as_completed(futures):
            result = future.result()
            filename = result['filename']
            file_path = os.path.join(TARGET_FOLDER, filename)

            if result['status'] == 'success':
                d = result['data']
                d["filename"] = filename
                if job_id:
                    d["job_id"] = job_id

                # Multi-factor score calculation
                d["match_score"] = calculate_multi_factor_score(d, expected_skills, jd_skills)

                with db_lock:
                    extracted_candidates_data.append(d)

                shutil.move(file_path, os.path.join(PROCESSED_FOLDER, filename))
                print(f"Processed: {filename}")
            else:
                logging.error(f"Extraction failed for {filename}: {result['error']}")
                shutil.move(file_path, os.path.join(FAILED_FOLDER, filename))
                print(f"Error: {filename}")

    with open(JSON_OUTPUT_FILE, 'w', encoding='utf-8') as json_file:
        json.dump(extracted_candidates_data, json_file, indent=4)

    print(f"Pipeline finished. {len(extracted_candidates_data)} candidates processed.")


if __name__ == "__main__":
    run_ats_pipeline()